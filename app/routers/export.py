import io
import json
from fastapi import APIRouter, HTTPException, Depends, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListItem, ListFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from app.db.database import get_db
from app.db.models import Summary, Quiz, ExamPrep, User
from app.routers.auth import get_current_user

router = APIRouter(prefix="/export", tags=["Export"])


def markdown_to_docx(title: str, text: str) -> io.BytesIO:
    """Convert basic markdown text to a DOCX document."""
    doc = DocxDocument()
    doc.add_heading(title, 0)

    for line in text.split("\n"):
        line_strip = line.strip()
        if not line_strip:
            continue
        
        # Headers
        if line_strip.startswith("# "):
            doc.add_heading(line_strip[2:], level=1)
        elif line_strip.startswith("## "):
            doc.add_heading(line_strip[3:], level=2)
        elif line_strip.startswith("### "):
            doc.add_heading(line_strip[4:], level=3)
        # Bullet list
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            doc.add_paragraph(line_strip[2:], style="List Bullet")
        # Numbered list
        elif line_strip[0].isdigit() and line_strip.split(".")[0].isdigit() and ". " in line_strip:
            doc.add_paragraph(line_strip.split(". ", 1)[1], style="List Number")
        else:
            # Basic bold formatting helper
            p = doc.add_paragraph()
            parts = line.split("**")
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def markdown_to_pdf(title: str, text: str) -> io.BytesIO:
    """Convert basic markdown text to a formatted PDF using reportlab."""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#1e3a5f'),
        spaceAfter=15
    )
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading2'],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0f2040'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=8
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        leftIndent=15,
        textColor=colors.HexColor('#333333'),
        spaceAfter=4
    )

    story = [Paragraph(title, title_style), Spacer(1, 10)]

    for line in text.split("\n"):
        line_strip = line.strip()
        if not line_strip:
            continue

        # HTML-like tag helper for reportlab formatting (replacing markdown ** with <b>)
        formatted_line = line
        parts = formatted_line.split("**")
        if len(parts) > 1:
            new_line = ""
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    new_line += f"<b>{part}</b>"
                else:
                    new_line += part
            formatted_line = new_line

        if line_strip.startswith("# "):
            story.append(Paragraph(formatted_line[2:], h1_style))
        elif line_strip.startswith("## ") or line_strip.startswith("### "):
            header_text = formatted_line.lstrip("#").strip()
            story.append(Paragraph(header_text, h1_style))
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            story.append(Paragraph(f"&bull; {formatted_line[2:]}", bullet_style))
        elif line_strip[0].isdigit() and line_strip.split(".")[0].isdigit() and ". " in line_strip:
            num, rest = formatted_line.split(". ", 1)
            story.append(Paragraph(f"{num}. {rest}", bullet_style))
        else:
            story.append(Paragraph(formatted_line, body_style))

    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer


def format_quiz_to_markdown(quiz_data: dict) -> str:
    """Format Quiz JSON data into a clean Markdown review sheet."""
    lines = []
    lines.append(f"# Practice Quiz: {quiz_data.get('question_type', 'General').upper()}")
    lines.append(f"**Difficulty Level:** {quiz_data.get('difficulty', 'Intermediate').upper()}")
    lines.append(f"**Number of Questions:** {len(quiz_data.get('questions', []))}")
    lines.append("\n---\n")

    solutions = []

    for i, q in enumerate(quiz_data.get("questions", []), 1):
        lines.append(f"### Question {i} ({q.get('type', 'mcq').upper()})")
        lines.append(q.get("question", ""))
        
        if q.get("options"):
            for key, val in q["options"].items():
                lines.append(f"- **{key}**: {val}")
        lines.append("")
        
        # Add to solutions section
        solutions.append(f"### Solution {i}")
        solutions.append(f"**Correct Answer:** {q.get('correct', '')}")
        solutions.append(f"**Explanation:** {q.get('explanation', '')}")
        if q.get("page_reference"):
            solutions.append(f"**Reference:** Page {q.get('page_reference')}")
        solutions.append("")

    lines.append("\n---\n")
    lines.append("# Answers & Explanations")
    lines.append("\n")
    lines.extend(solutions)

    return "\n".join(lines)


@router.get("/{feature_type}/{item_id}")
async def export_item(
    feature_type: str,  # "summary", "quiz", "examprep"
    item_id: int,
    format: str = "markdown",  # "pdf", "docx", "markdown"
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export study items (summaries, quizzes, exam preparation sheets)
    into formatted PDF, Microsoft Word (DOCX), or plain Markdown files.
    """
    if feature_type not in ["summary", "quiz", "examprep"]:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Invalid feature type for export")
    if format not in ["pdf", "docx", "markdown"]:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Invalid export format")

    content_text = ""
    title = ""

    # Fetch content from SQLite based on user permissions
    if feature_type == "summary":
        item = db.query(Summary).filter(Summary.id == item_id, Summary.user_id == current_user.id).first()
        if not item:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Summary not found")
        title = f"Document Summary - {item.level.upper()}"
        content_text = item.content

    elif feature_type == "examprep":
        item = db.query(ExamPrep).filter(ExamPrep.id == item_id, ExamPrep.user_id == current_user.id).first()
        if not item:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Exam preparation guide not found")
        title = f"Exam Preparation - {item.prep_type.replace('_', ' ').title()}"
        content_text = item.content

    elif feature_type == "quiz":
        item = db.query(Quiz).filter(Quiz.id == item_id, Quiz.user_id == current_user.id).first()
        if not item:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Quiz not found")
        try:
            quiz_data = json.loads(item.content)
            # Add metadata keys if missing
            quiz_data["difficulty"] = item.difficulty
            quiz_data["question_type"] = item.question_type
            title = f"Quiz Sheet ({item.difficulty.upper()})"
            content_text = format_quiz_to_markdown(quiz_data)
        except Exception:
            # Fallback if content was raw string
            title = "Quiz Sheet"
            content_text = item.content

    # Export formats
    if format == "markdown":
        file_io = io.BytesIO(content_text.encode("utf-8"))
        return StreamingResponse(
            file_io,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={feature_type}_{item_id}.md"}
        )

    elif format == "docx":
        file_io = markdown_to_docx(title, content_text)
        return StreamingResponse(
            file_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={feature_type}_{item_id}.docx"}
        )

    elif format == "pdf":
        file_io = markdown_to_pdf(title, content_text)
        return StreamingResponse(
            file_io,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={feature_type}_{item_id}.pdf"}
        )
